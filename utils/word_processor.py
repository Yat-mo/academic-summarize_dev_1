from docx import Document
import tempfile
import config

class WordProcessor:
    def __init__(self):
        pass

    def extract_text(self, file):
        """
        从Word文档中提取文本内容
        :param file: Word文档文件对象
        :return: 提取的文本内容列表，每个元素都是一个文本块
        """
        # 创建临时文件来保存上传的文件内容
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_file_path = tmp_file.name

        # 使用python-docx打开文档
        doc = Document(tmp_file_path)
        
        # 提取文本
        text_parts = []
        current_chunk = ""
        
        # 提取正文段落
        for para in doc.paragraphs:
            if para.text.strip():  # 只添加非空段落
                words = para.text.split()
                for word in words:
                    if len(current_chunk) + len(word) + 1 > config.CHUNK_SIZE:
                        if current_chunk:
                            text_parts.append(current_chunk.strip())
                        current_chunk = word
                    else:
                        current_chunk += " " + word if current_chunk else word
                
                # 每个段落结束时添加换行
                if current_chunk:
                    current_chunk += "\n"
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # 只添加非空单元格
                        row_text.append(cell.text.strip())
                
                if row_text:
                    row_content = " | ".join(row_text)
                    words = row_content.split()
                    for word in words:
                        if len(current_chunk) + len(word) + 1 > config.CHUNK_SIZE:
                            if current_chunk:
                                text_parts.append(current_chunk.strip())
                            current_chunk = word
                        else:
                            current_chunk += " " + word if current_chunk else word
                    
                    # 每行结束时添加换行
                    if current_chunk:
                        current_chunk += "\n"
        
        # 添加最后一个块
        if current_chunk:
            text_parts.append(current_chunk.strip())
        
        # 删除临时文件
        import os
        os.unlink(tmp_file_path)
        
        return text_parts if text_parts else ["文档为空"]
